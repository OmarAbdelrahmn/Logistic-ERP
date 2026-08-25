from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path(r"C:\Users\omarf\OneDrive\Documents\ChatGPT\Logistic ERP")
FRONTS = WORKSPACE / "tmp" / "amazon_contact_sheet" / "fronts"
OUTPUT = WORKSPACE / "output" / "Amazon_front_pages_6_per_page.docx"

# Compact reference guide with a named image-grid override:
# Letter landscape, 0.4-in margins, no table borders, 3.2-in image width.
PAGE_WIDTH = Inches(11)
PAGE_HEIGHT = Inches(8.5)
MARGIN = Inches(0.4)
CELL_WIDTH_DXA = 4896
TABLE_WIDTH_DXA = 14688


def set_cell_margins(cell, top=60, start=60, bottom=60, end=60):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "nil")

    grid = table._tbl.tblGrid
    for grid_col in grid.gridCol_lst:
        grid_col.set(qn("w:w"), str(CELL_WIDTH_DXA))
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            set_cell_width(cell, CELL_WIDTH_DXA)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_sheet(doc, images):
    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table)
    for cell, image in zip((c for r in table.rows for c in r.cells), images):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1)
        picture = paragraph.add_run().add_picture(str(image), width=Inches(3.2))
        # Accessible description based on the original source-file ID.
        doc_pr = picture._inline.docPr
        source_id = image.stem.split("_", 1)[1].rsplit("-", 1)[0]
        doc_pr.set("descr", f"Front page from Amazon file {source_id}")
        caption = cell.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(0)
        caption.paragraph_format.space_after = Pt(0)
        run = caption.add_run(source_id)
        set_run_font(run, 8, "555555")


def main():
    images = sorted(FRONTS.glob("*.png"))
    if len(images) != 24:
        raise ValueError(f"Expected 24 first-page images, found {len(images)}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for page in range(4):
        add_sheet(doc, images[page * 6 : (page + 1) * 6])
        if page < 3:
            doc.add_page_break()
    doc.core_properties.title = "Amazon Front Pages"
    doc.core_properties.subject = "First page from each supplied Amazon PDF"
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
