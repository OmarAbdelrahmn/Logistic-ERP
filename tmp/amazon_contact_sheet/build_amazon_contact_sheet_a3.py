from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches


WORKSPACE = Path(r"C:\Users\omarf\OneDrive\Documents\ChatGPT\Logistic ERP")
FRONTS = WORKSPACE / "tmp" / "amazon_contact_sheet" / "fronts"
SHEETS = WORKSPACE / "tmp" / "amazon_contact_sheet" / "a3_sheets"
OUTPUT = WORKSPACE / "output" / "Amazon_front_pages_6_per_page.docx"

# A3 landscape sheet, 3 x 2 image layout.  This rotates the page layout and
# increases each readable, upright card from 3.6 in to nearly 5.0 in wide.
CANVAS_SIZE = (4960, 3000)
CARD_WIDTH = 1560
CARD_HEIGHT = 987
LEFT = 90
TOP = 420
COLUMN_GAP = 100
ROW_GAP = 120
LABEL_GAP = 20


def label_font():
    for path in (Path(r"C:\Windows\Fonts\calibri.ttf"), Path(r"C:\Windows\Fonts\arial.ttf")):
        if path.exists():
            return ImageFont.truetype(str(path), 44)
    return ImageFont.load_default()


def make_sheet(index, images):
    canvas = Image.new("RGB", CANVAS_SIZE, "white")
    draw = ImageDraw.Draw(canvas)
    font = label_font()
    for position, image_path in enumerate(images):
        row, col = divmod(position, 3)
        x = LEFT + col * (CARD_WIDTH + COLUMN_GAP)
        y = TOP + row * (CARD_HEIGHT + LABEL_GAP + 55 + ROW_GAP)
        with Image.open(image_path) as source:
            card = source.convert("RGB").resize((CARD_WIDTH, CARD_HEIGHT), Image.Resampling.LANCZOS)
        canvas.paste(card, (x, y))
        source_id = image_path.stem.split("_", 1)[1].rsplit("-", 1)[0]
        box = draw.textbbox((0, 0), source_id, font=font)
        draw.text((x + (CARD_WIDTH - (box[2] - box[0])) // 2, y + CARD_HEIGHT + LABEL_GAP), source_id, font=font, fill="#555555")
    target = SHEETS / f"sheet_{index:02d}.png"
    canvas.save(target, optimize=True)
    return target


def configure_document(doc, section=None):
    section = section or doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(16.54)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def main():
    images = sorted(FRONTS.glob("*.png"))
    if len(images) != 24:
        raise ValueError(f"Expected 24 first-page images, found {len(images)}")
    SHEETS.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    sheets = [make_sheet(page + 1, images[page * 6 : (page + 1) * 6]) for page in range(4)]
    doc = Document()
    configure_document(doc)
    for index, sheet in enumerate(sheets):
        if index:
            configure_document(doc, doc.add_section(WD_SECTION.NEW_PAGE))
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Inches(0)
        paragraph.paragraph_format.space_after = Inches(0)
        picture = paragraph.add_run().add_picture(str(sheet), width=Inches(15.74))
        picture._inline.docPr.set("descr", f"Large-format Amazon first-page contact sheet {index + 1} of 4")
    doc.core_properties.title = "Amazon Front Pages - Large Format"
    doc.core_properties.subject = "First page from each supplied Amazon PDF"
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
