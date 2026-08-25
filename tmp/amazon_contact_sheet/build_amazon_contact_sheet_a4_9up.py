from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches


WORKSPACE = Path(r"C:\Users\omarf\OneDrive\Documents\ChatGPT\Logistic ERP")
FRONTS = WORKSPACE / "tmp" / "amazon_contact_sheet" / "fronts"
SHEETS = WORKSPACE / "tmp" / "amazon_contact_sheet" / "a4_9up_sheets"
OUTPUT = WORKSPACE / "output" / "Amazon_front_pages_9_per_page.docx"

# A4 landscape at 300 dpi.  The contact sheet itself has no outer margins and
# includes only the source card images (no file-name captions).
CANVAS_SIZE = (3508, 2446)
CARD_WIDTH = 1136
CARD_HEIGHT = 719
H_GAP = 50
V_GAP = 144


def make_sheet(index, images):
    canvas = Image.new("RGB", CANVAS_SIZE, "white")
    for position, image_path in enumerate(images):
        row, col = divmod(position, 3)
        x = col * (CARD_WIDTH + H_GAP)
        y = row * (CARD_HEIGHT + V_GAP)
        with Image.open(image_path) as source:
            card = source.convert("RGB").resize((CARD_WIDTH, CARD_HEIGHT), Image.Resampling.LANCZOS)
        canvas.paste(card, (x, y))
    target = SHEETS / f"sheet_{index:02d}.png"
    canvas.save(target, optimize=True)
    return target


def configure_document(doc, section=None):
    section = section or doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    safe_margin = Inches(0.19685)  # 0.5 cm
    section.top_margin = safe_margin
    section.bottom_margin = safe_margin
    section.left_margin = safe_margin
    section.right_margin = safe_margin
    section.header_distance = safe_margin
    section.footer_distance = safe_margin
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def main():
    images = sorted(FRONTS.glob("*.png"))
    if len(images) != 25:
        raise ValueError(f"Expected 25 first-page images, found {len(images)}")
    SHEETS.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    sheets = [make_sheet(page + 1, images[page * 9 : (page + 1) * 9]) for page in range(3)]
    doc = Document()
    configure_document(doc)
    for index, sheet in enumerate(sheets):
        if index:
            configure_document(doc, doc.add_section(WD_SECTION.NEW_PAGE))
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Inches(0)
        paragraph.paragraph_format.space_after = Inches(0)
        paragraph.paragraph_format.left_indent = Inches(0)
        paragraph.paragraph_format.right_indent = Inches(0)
        picture = paragraph.add_run().add_picture(str(sheet), width=Inches(11.2963))
        picture._inline.docPr.set("descr", f"A4 Amazon card contact sheet {index + 1} of 3")
    doc.core_properties.title = "Amazon Front Pages - A4 9-up"
    doc.core_properties.subject = "First page from each supplied Amazon PDF"
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
